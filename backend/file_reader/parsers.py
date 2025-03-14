"""
File parsers module for handling different file types.
This module provides specialized parsers for various file formats including:
- Text files (TXT, CSV, YAML, JSON, XML)
- Document files (PDF, DOCX)
- Spreadsheets (XLSX, ODS)
- Code/config files (PY, JS, TOML, etc.)
- Vector/CAD files (DWG, DXF, SVG)
- Archives (ZIP, TAR, RAR, 7Z)
"""

import ast
import asyncio
import csv
import json
import logging
import tarfile
import xml.etree.ElementTree as ET
import zipfile
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Dict, Optional

import ezdxf
import pandas as pd
import py7zr
import rarfile  # RAR support
import tomli
import yaml
from docx import Document
from pypdf import PdfReader

logger = logging.getLogger(__name__)


class BaseParser(ABC):
    """Abstract base class for file parsers."""

    @abstractmethod
    async def parse(self, file_path: Path) -> Dict[str, Any]:
        """
        Parse a file asynchronously and return its structured data.

        :param file_path: Path to the file to parse
        :return: A dictionary containing parsed data (keys may differ by parser).
        """
        pass

    @abstractmethod
    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        """
        Generate a textual preview for the parsed data.

        :param parsed_data: The dictionary returned by `parse()`
        :param max_length: Maximum length of the preview string
        :return: A string preview of the data
        """
        pass


class TextParser(BaseParser):
    """Parser for basic text files (TXT, MD, etc.)."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        from backend.utils.async_io import AsyncFileIO

        try:
            content = await AsyncFileIO.read(file_path)  # Async read
            if content.startswith("# "):
                # Means either file was ignored or an error occurred
                return {
                    "error": content,
                    "content": "",
                    "encoding": "",
                    "line_count": 0,
                }

            encoding_used = "utf-8"
            line_count = content.count("\n") + 1
            return {
                "content": content,
                "encoding": encoding_used,
                "line_count": line_count,
            }
        except UnicodeDecodeError:
            # Skipping fallback for brevity; you could do another read with 'latin-1'
            return {
                "error": "UnicodeDecodeError on text file",
                "content": "",
                "encoding": "",
                "line_count": 0,
            }

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        content = parsed_data.get("content", "")
        return content[:max_length] + "..." if len(content) > max_length else content


class JsonParser(BaseParser):
    """Parser for JSON files."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        from backend.utils.async_io import AsyncFileIO

        json_data = await AsyncFileIO.read_json(file_path)
        return {
            "content": json_data,
            "structure": type(json_data).__name__,
            "size": len(json.dumps(json_data)),
        }

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        preview = json.dumps(parsed_data["content"], indent=2)
        return preview[:max_length] + "..." if len(preview) > max_length else preview


class YamlParser(BaseParser):
    """Parser for YAML files."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        from backend.utils.async_io import AsyncFileIO

        content = await AsyncFileIO.read(file_path)
        data = yaml.safe_load(content)
        return {
            "content": data,
            "preview": content[:1000] if len(content) > 1000 else content,
            "metadata": {
                "keys": list(data.keys()) if isinstance(data, dict) else None,
                "size": len(data) if isinstance(data, (dict, list)) else None,
            },
        }

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        preview = yaml.dump(parsed_data["content"], default_flow_style=False)
        return preview[:max_length] + "..." if len(preview) > max_length else preview


class CsvParser(BaseParser):
    """Parser for CSV files."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        from aichemist_codex.utils import AsyncFileIO

        try:
            content = await AsyncFileIO.read(file_path)  # Async read entire CSV
            if content.startswith("# "):
                return {
                    "error": content,
                    "header": None,
                    "rows": [],
                    "row_count": 0,
                    "column_count": 0,
                }

            rows = []
            reader = csv.reader(content.splitlines())
            header = next(reader, None)
            for row in reader:
                rows.append(row)

            return {
                "header": header,
                "rows": rows,
                "row_count": len(rows),
                "column_count": len(header) if header else 0,
            }
        except Exception as e:
            logger.error(f"Error parsing CSV: {e}")
            return {
                "error": str(e),
                "header": None,
                "rows": [],
                "row_count": 0,
                "column_count": 0,
            }

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        preview_lines = []
        if parsed_data["header"]:
            preview_lines.append(",".join(parsed_data["header"]))

        for row in parsed_data["rows"][:5]:
            preview_lines.append(",".join(row))

        preview = "\n".join(preview_lines)
        return preview[:max_length] + "..." if len(preview) > max_length else preview


class XmlParser(BaseParser):
    """Parser for XML files."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        import xml.etree.ElementTree as ET

        from aichemist_codex.utils import AsyncFileIO

        content = await AsyncFileIO.read(file_path)
        if content.startswith("# "):
            return {
                "error": content,
                "content": "",
                "metadata": {},
                "preview": "",
            }
        root = ET.fromstring(content)
        return {
            "content": content,
            "preview": content[:1000] if len(content) > 1000 else content,
            "metadata": {
                "root_tag": root.tag,
                "children_count": len(list(root)),
                "attributes": dict(root.attrib),
            },
        }

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        preview = f"Root: {parsed_data['metadata']['root_tag']}\n"
        content_preview = parsed_data["content"]
        if len(preview) + len(content_preview) > max_length:
            return preview + content_preview[: max_length - len(preview)] + "..."
        return preview + content_preview


class DocumentParser(BaseParser):
    """Parser for document files (PDF, DOCX, etc.)."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:

        suffix = file_path.suffix.lower()
        try:
            if suffix == ".pdf":
                return await self._parse_pdf(file_path)
            elif suffix == ".docx":
                return await self._parse_docx(file_path)
            else:
                raise ValueError(f"Unsupported document format: {suffix}")
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {e}")
            raise

    async def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        loop = asyncio.get_running_loop()

        def parse_pdf_sync(p):
            reader = PdfReader(str(p))
            text_content = []
            for page in reader.pages:
                text_content.append(page.extract_text())
            return {
                "content": "\n".join(text_content),
                "metadata": reader.metadata,
                "pages": len(reader.pages),
            }

        return await loop.run_in_executor(None, parse_pdf_sync, file_path)

    async def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        import asyncio

        loop = asyncio.get_running_loop()

        def parse_docx_sync(p):
            doc = Document(str(p))
            content = "\n".join(para.text for para in doc.paragraphs)
            return {
                "content": content,
                "metadata": {
                    "sections": len(doc.sections),
                    "paragraphs": len(doc.paragraphs),
                },
            }

        return await loop.run_in_executor(None, parse_docx_sync, file_path)

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        content = parsed_data.get("content", "")
        return content[:max_length] + "..." if len(content) > max_length else content


class SpreadsheetParser(BaseParser):
    """Parser for spreadsheet files (CSV, XLSX, ODS)."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        suffix = file_path.suffix.lower()
        try:
            if suffix == ".csv":
                return await self._parse_csv(file_path)
            elif suffix == ".xlsx":
                return await self._parse_xlsx(file_path)
            elif suffix == ".ods":
                return await self._parse_ods(file_path)
            else:
                raise ValueError(f"Unsupported spreadsheet format: {suffix}")
        except Exception as e:
            logger.error(f"Error parsing spreadsheet {file_path}: {str(e)}")
            raise

    async def _parse_csv(self, file_path: Path) -> Dict[str, Any]:
        import asyncio

        def parse_csv_sync(p):
            df = pd.read_csv(p)
            preview = df.head().to_string()
            return {
                "content": df.to_dict(),
                "preview": preview,
                "metadata": {
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": df.columns.tolist(),
                },
            }

        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, parse_csv_sync, file_path)

    async def _parse_xlsx(self, file_path: Path) -> Dict[str, Any]:
        import asyncio

        def parse_xlsx_sync(p):
            df = pd.read_excel(p)
            preview = df.head().to_string()
            return {
                "content": df.to_dict(),
                "preview": preview,
                "metadata": {
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": df.columns.tolist(),
                    "sheet_names": pd.ExcelFile(str(p)).sheet_names,
                },
            }

        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, parse_xlsx_sync, file_path)

    async def _parse_ods(self, file_path: Path) -> Dict[str, Any]:
        import asyncio

        def parse_ods_sync(p):
            df = pd.read_excel(p, engine="odf")
            preview = df.head().to_string()
            return {
                "content": df.to_dict(),
                "preview": preview,
                "metadata": {
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": df.columns.tolist(),
                },
            }

        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, parse_ods_sync, file_path)

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        preview = parsed_data.get("preview", "")
        return preview[:max_length] + "..." if len(preview) > max_length else preview


class CodeParser(BaseParser):
    """Parser for code and configuration files (Python, JS, JSON, YAML, XML, TOML)."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        suffix = file_path.suffix.lower()
        try:
            if suffix == ".py":
                return await self._parse_python(file_path)
            elif suffix == ".js":
                return await self._parse_javascript(file_path)
            elif suffix == ".json":
                return await self._parse_json(file_path)
            elif suffix in [".yaml", ".yml"]:
                return await self._parse_yaml(file_path)
            elif suffix == ".xml":
                return await self._parse_xml(file_path)
            elif suffix == ".toml":
                return await self._parse_toml(file_path)
            else:
                raise ValueError(f"Unsupported code/config format: {suffix}")
        except Exception as e:
            logger.error(f"Error parsing code/config file {file_path}: {str(e)}")
            raise

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        preview = parsed_data.get("preview", "")
        return preview[:max_length] + "..." if len(preview) > max_length else preview

    async def _parse_python(self, file_path: Path) -> Dict[str, Any]:
        from aichemist_codex.utils import AsyncFileIO

        content = await AsyncFileIO.read(file_path)
        if content.startswith("# "):
            return {"error": content, "preview": content, "metadata": {}}

        try:
            tree = ast.parse(content)
            classes = [n.name for n in ast.walk(tree) if isinstance(n, ast.ClassDef)]
            functions = [
                n.name for n in ast.walk(tree) if isinstance(n, ast.FunctionDef)
            ]
            imports = [
                n for n in ast.walk(tree) if isinstance(n, (ast.Import, ast.ImportFrom))
            ]
            return {
                "content": content,
                "preview": content[:1000] if len(content) > 1000 else content,
                "metadata": {
                    "classes": classes,
                    "functions": functions,
                    "import_count": len(imports),
                    "loc": len(content.splitlines()),
                },
            }
        except Exception as e:
            logger.error(f"Python parsing error: {str(e)}")
            raise

    async def _parse_javascript(self, file_path: Path) -> Dict[str, Any]:
        from aichemist_codex.utils import AsyncFileIO

        content = await AsyncFileIO.read(file_path)
        if content.startswith("// "):
            return {"error": content, "preview": content, "metadata": {}}

        return {
            "content": content,
            "preview": content[:1000] if len(content) > 1000 else content,
            "metadata": {"loc": len(content.splitlines())},
        }

    async def _parse_json(self, file_path: Path) -> Dict[str, Any]:
        from aichemist_codex.utils import AsyncFileIO

        json_data = await AsyncFileIO.read_json(file_path)
        if not json_data:
            return {
                "error": "Failed to parse JSON file",
                "preview": "Error reading JSON",
                "metadata": {},
            }
        content = await AsyncFileIO.read(file_path)
        return {
            "content": json_data,
            "preview": content[:1000] if len(content) > 1000 else content,
            "metadata": {
                "keys": list(json_data.keys()) if isinstance(json_data, dict) else None,
                "size": (
                    len(json_data) if isinstance(json_data, (dict, list)) else None
                ),
            },
        }

    async def _parse_yaml(self, file_path: Path) -> Dict[str, Any]:
        from aichemist_codex.utils import AsyncFileIO

        content = await AsyncFileIO.read(file_path)
        data = yaml.safe_load(content)
        return {
            "content": data,
            "preview": content[:1000] if len(content) > 1000 else content,
            "metadata": {
                "keys": list(data.keys()) if isinstance(data, dict) else None,
                "size": len(data) if isinstance(data, (dict, list)) else None,
            },
        }

    async def _parse_xml(self, file_path: Path) -> Dict[str, Any]:
        from aichemist_codex.utils import AsyncFileIO

        content = await AsyncFileIO.read(file_path)
        root = ET.fromstring(content)
        return {
            "content": content,
            "preview": content[:1000] if len(content) > 1000 else content,
            "metadata": {
                "root_tag": root.tag,
                "children_count": len(list(root)),
                "attributes": dict(root.attrib),
            },
        }

    async def _parse_toml(self, file_path: Path) -> Dict[str, Any]:
        from aichemist_codex.utils import AsyncFileIO

        content = await AsyncFileIO.read_binary(file_path)
        try:
            data = tomli.loads(content.decode("utf-8"))
            return {
                "content": data,
                "preview": str(data)[:1000],
                "metadata": {
                    "keys": list(data.keys()) if isinstance(data, dict) else None,
                    "size": len(data) if isinstance(data, (dict, list)) else None,
                },
            }
        except Exception as e:
            logger.error(f"TOML parsing error: {str(e)}")
            raise


class VectorParser(BaseParser):
    """Parser for CAD and vector files (DWG, DXF, SVG)."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        suffix = file_path.suffix.lower()
        try:
            if suffix in [".dwg", ".dxf"]:
                return await self._parse_cad(file_path)
            elif suffix == ".svg":
                return await self._parse_svg(file_path)
            else:
                raise ValueError(f"Unsupported vector format: {suffix}")
        except Exception as e:
            logger.error(f"Error parsing vector file {file_path}: {str(e)}")
            raise

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        preview = parsed_data.get("preview", "")
        return preview[:max_length] + "..." if len(preview) > max_length else preview

    async def _parse_cad(self, file_path: Path) -> Dict[str, Any]:
        try:
            doc = ezdxf.readfile(str(file_path))
            modelspace = doc.modelspace()
            entities = {
                "lines": len(modelspace.query("LINE")),
                "circles": len(modelspace.query("CIRCLE")),
                "arcs": len(modelspace.query("ARC")),
                "polylines": len(modelspace.query("LWPOLYLINE")),
                "text": len(modelspace.query("TEXT")),
            }
            layers = [layer.dxf.name for layer in doc.layers]
            metadata = {
                "filename": file_path.name,
                "created_by": doc.header["$TDCREATE"],
                "last_modified": doc.header["$TDUPDATE"],
                "drawing_units": doc.header["$INSUNITS"],
                "layers": layers,
                "entity_counts": entities,
            }
            return {
                "content": str(doc.entitydb),
                "preview": f"CAD drawing with {sum(entities.values())} entities across {len(layers)} layers",
                "metadata": metadata,
            }
        except Exception as e:
            logger.error(f"CAD parsing error: {str(e)}")
            raise

    async def _parse_svg(self, file_path: Path) -> Dict[str, Any]:
        from aichemist_codex.utils import AsyncFileIO

        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            width = root.get("width", "unknown")
            height = root.get("height", "unknown")
            viewBox = root.get("viewBox", "unknown")
            elements = {
                "path": len(root.findall(".//{*}path")),
                "rect": len(root.findall(".//{*}rect")),
                "circle": len(root.findall(".//{*}circle")),
                "text": len(root.findall(".//{*}text")),
                "group": len(root.findall(".//{*}g")),
            }
            content = await AsyncFileIO.read(file_path)
            if content.startswith("# "):
                return {"error": content, "preview": content, "metadata": {}}
            return {
                "content": content,
                "preview": f"SVG image ({width}x{height}) with {sum(elements.values())} elements",
                "metadata": {
                    "dimensions": {
                        "width": width,
                        "height": height,
                        "viewBox": viewBox,
                    },
                    "element_counts": elements,
                    "xmlns": root.get("xmlns", "unknown"),
                    "version": root.get("version", "unknown"),
                },
            }
        except Exception as e:
            logger.error(f"SVG parsing error: {str(e)}")
            raise


class ArchiveParser(BaseParser):
    """Parser for archive files (ZIP, TAR, RAR, 7Z)."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        from aichemist_codex.utils import AsyncFileIO

        try:
            if not await AsyncFileIO.exists(file_path):
                raise FileNotFoundError(f"Archive file not found: {file_path}")

            suffix = file_path.suffix.lower()
            files_list = []

            if suffix == ".zip":
                with zipfile.ZipFile(file_path, "r") as archive:
                    files_list = archive.namelist()
            elif suffix in [".tar", ".tgz", ".gz", ".bz2"]:
                with tarfile.open(file_path, "r") as archive:
                    files_list = archive.getnames()
            elif suffix == ".rar":
                try:
                    with rarfile.RarFile(file_path, "r") as archive:
                        files_list = archive.namelist()
                except Exception as e:
                    if isinstance(e, rarfile.NotRarFile):
                        raise ValueError(f"Unsupported archive format: {suffix}") from e
                    else:
                        raise
            elif suffix == ".7z":
                with py7zr.SevenZipFile(file_path, "r") as archive:
                    files_list = archive.getnames()
            else:
                raise ValueError(f"Unsupported archive format: {suffix}")

            return {"files": files_list, "count": len(files_list)}
        except Exception as e:
            logger.error(f"Archive parsing failed for {file_path}: {e}", exc_info=True)
            raise

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        files = parsed_data.get("files", [])
        preview = "\n".join(files)
        return preview[:max_length] + "..." if len(preview) > max_length else preview


def get_parser_for_mime_type(mime_type: str) -> Optional[BaseParser]:
    """
    Factory function to get the appropriate parser for a MIME type.
    """
    parsers = {
        "text/plain": TextParser(),
        "text/markdown": TextParser(),
        "text/html": TextParser(),
        "application/json": JsonParser(),
        "application/yaml": YamlParser(),
        "text/yaml": YamlParser(),
        "text/csv": CsvParser(),
        "application/xml": XmlParser(),
        "text/xml": XmlParser(),
        "application/pdf": DocumentParser(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DocumentParser(),
        "application/vnd.oasis.opendocument.text": DocumentParser(),
        "application/epub+zip": DocumentParser(),
        "application/zip": ArchiveParser(),
        "application/x-tar": ArchiveParser(),
        "application/x-rar-compressed": ArchiveParser(),
        "application/x-7z-compressed": ArchiveParser(),
        "application/gzip": ArchiveParser(),
        "application/x-bzip2": ArchiveParser(),
        "image/vnd.dxf": VectorParser(),
        "image/x-dwg": VectorParser(),
        "image/svg+xml": VectorParser(),
        "text/x-python": CodeParser(),
        "application/javascript": CodeParser(),
        "application/toml": CodeParser(),
    }

    if mime_type in parsers:
        return parsers[mime_type]

    main_type = mime_type.split("/")[0]
    if main_type == "text":
        return TextParser()

    return None
