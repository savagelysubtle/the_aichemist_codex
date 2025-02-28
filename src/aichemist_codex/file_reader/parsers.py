"""
File parsers module for handling different file types.

This module provides specialized parsers for various file formats including:
- Text-based files (TXT, MD, LOG, RTF, HTML)
- Document files (PDF, DOCX, ODT, EPUB)
- Spreadsheets (CSV, XLSX, ODS)
- Code & configuration files (PY, JS, JSON, YAML, XML, TOML)
- CAD & vector files (DWG, DXF, SVG)
- Archives (ZIP, TAR, RAR, 7Z)
"""

import ast
import csv
import json
import logging
import tarfile
import xml.etree.ElementTree as ET
import zipfile
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Dict, Optional

import aiofiles
import ezdxf
import pandas as pd
import py7zr
import rarfile  # New import for RAR support
import tomli
import yaml
from docx import Document
from pypdf import PdfReader

logger = logging.getLogger(__name__)


class BaseParser(ABC):
    """Base class for all file parsers."""

    @abstractmethod
    async def parse(self, file_path: Path) -> Dict[str, Any]:
        """
        Parse the file and return its contents and metadata.

        Args:
            file_path: Path to the file to parse

        Returns:
            Dict containing parsed content and metadata
        """
        pass

    @abstractmethod
    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        """
        Generate a preview of the parsed content.

        Args:
            parsed_data: The parsed file data
            max_length: Maximum length of the preview

        Returns:
            A string preview of the content
        """
        pass


class TextParser(BaseParser):
    """Parser for text-based files (TXT, MD, LOG, etc.)."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            return {
                "content": content,
                "encoding": "utf-8",
                "line_count": content.count("\n") + 1,
            }
        except UnicodeDecodeError:
            # Try alternative encoding
            with open(file_path, "r", encoding="latin-1") as f:
                content = f.read()
            return {
                "content": content,
                "encoding": "latin-1",
                "line_count": content.count("\n") + 1,
            }

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        content = parsed_data["content"]
        if len(content) > max_length:
            return content[:max_length] + "..."
        return content


class JsonParser(BaseParser):
    """Parser for JSON files."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        with open(file_path, "r", encoding="utf-8") as f:
            content = json.load(f)
        return {
            "content": content,
            "structure": type(content).__name__,
            "size": len(json.dumps(content)),
        }

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        preview = json.dumps(parsed_data["content"], indent=2)
        if len(preview) > max_length:
            return preview[:max_length] + "..."
        return preview


class YamlParser(BaseParser):
    """Parser for YAML files."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        with open(file_path, "r", encoding="utf-8") as f:
            content = yaml.safe_load(f)
        return {
            "content": content,
            "structure": type(content).__name__,
        }

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        preview = yaml.dump(parsed_data["content"], default_flow_style=False)
        if len(preview) > max_length:
            return preview[:max_length] + "..."
        return preview


class CsvParser(BaseParser):
    """Parser for CSV files."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        rows = []
        with open(file_path, "r", encoding="utf-8", newline="") as f:
            reader = csv.reader(f)
            header = next(reader, None)
            for row in reader:
                rows.append(row)

        return {
            "header": header,
            "rows": rows,
            "row_count": len(rows),
            "column_count": len(header) if header else 0,
        }

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        preview_lines = []
        if parsed_data["header"]:
            preview_lines.append(",".join(parsed_data["header"]))

        for row in parsed_data["rows"][:5]:  # Show first 5 rows
            preview_lines.append(",".join(row))

        preview = "\n".join(preview_lines)
        if len(preview) > max_length:
            return preview[:max_length] + "..."
        return preview


class XmlParser(BaseParser):
    """Parser for XML files."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        tree = ET.parse(file_path)
        root = tree.getroot()

        def element_to_dict(element):
            result = {}
            for child in element:
                if len(child) > 0:
                    result[child.tag] = element_to_dict(child)
                else:
                    result[child.tag] = child.text
            return result

        return {
            "root_tag": root.tag,
            "content": element_to_dict(root),
            "attributes": dict(root.attrib),
        }

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        preview = f"Root: {parsed_data['root_tag']}\n"
        preview += f"Attributes: {json.dumps(parsed_data['attributes'], indent=2)}\n"
        content_preview = json.dumps(parsed_data["content"], indent=2)

        if len(preview) + len(content_preview) > max_length:
            return preview + content_preview[: max_length - len(preview)] + "..."
        return preview + content_preview


class DocumentParser(BaseParser):
    """Parser for document files (PDF, DOCX, ODT, EPUB)."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        """Parse document file and return its contents and metadata.

        Args:
            file_path: Path to the document file

        Returns:
            Dict containing parsed content and metadata

        Raises:
            ValueError: If the file format is not supported
        """
        suffix = file_path.suffix.lower()
        try:
            if suffix == ".pdf":
                return await self._parse_pdf(file_path)
            elif suffix == ".docx":
                return await self._parse_docx(file_path)
            elif suffix == ".odt":
                return await self._parse_odt(file_path)
            elif suffix == ".epub":
                return await self._parse_epub(file_path)
            else:
                raise ValueError(f"Unsupported document format: {suffix}")
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {str(e)}")
            raise

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        """Generate a preview of the document content.

        Args:
            parsed_data: The parsed document data
            max_length: Maximum length of the preview

        Returns:
            A string preview of the content
        """
        content = parsed_data.get("content", "")
        if len(content) > max_length:
            return content[:max_length] + "..."
        return content

    async def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF files using pypdf."""
        try:
            # Open file directly instead of using bytes
            reader = PdfReader(str(file_path))
            text_content = []
            for page in reader.pages:
                text_content.append(page.extract_text())

            return {
                "content": "\n".join(text_content),
                "metadata": reader.metadata,
                "pages": len(reader.pages),
            }
        except Exception as e:
            logger.error(f"PDF parsing error: {str(e)}")
            raise

    async def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX files."""
        try:
            doc = Document(file_path)
            content = []
            for para in doc.paragraphs:
                content.append(para.text)

            return {
                "content": "\n".join(content),
                "metadata": {
                    "sections": len(doc.sections),
                    "paragraphs": len(doc.paragraphs),
                },
            }
        except Exception as e:
            logger.error(f"DOCX parsing error: {str(e)}")
            raise


class SpreadsheetParser(BaseParser):
    """Parser for spreadsheet files (CSV, XLSX, ODS)."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        """Parse spreadsheet file and return its contents and metadata.

        Args:
            file_path: Path to the spreadsheet file

        Returns:
            Dict containing parsed content and metadata

        Raises:
            ValueError: If the file format is not supported
        """
        suffix = file_path.suffix.lower()
        try:
            if suffix == ".csv":
                return await self._parse_csv(file_path)
            elif suffix == ".xlsx":
                return await self._parse_xlsx(file_path)
            elif suffix == ".ods":
                return await self._parse_ods(file_path)
            else:
                raise ValueError(f"Unsupported spreadsheet format: {suffix}")
        except Exception as e:
            logger.error(f"Error parsing spreadsheet {file_path}: {str(e)}")
            raise

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        """Generate a preview of the spreadsheet content.

        Args:
            parsed_data: The parsed spreadsheet data
            max_length: Maximum length of the preview

        Returns:
            A string preview of the content
        """
        preview = parsed_data.get("preview", "")
        if len(preview) > max_length:
            return preview[:max_length] + "..."
        return preview

    async def _parse_csv(self, file_path: Path) -> Dict[str, Any]:
        """Parse CSV files."""
        try:
            df = pd.read_csv(file_path)
            preview = df.head().to_string()
            return {
                "content": df.to_dict(),
                "preview": preview,
                "metadata": {
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": df.columns.tolist(),
                },
            }
        except Exception as e:
            logger.error(f"CSV parsing error: {str(e)}")
            raise

    async def _parse_xlsx(self, file_path: Path) -> Dict[str, Any]:
        """Parse XLSX files."""
        try:
            df = pd.read_excel(file_path)
            preview = df.head().to_string()
            return {
                "content": df.to_dict(),
                "preview": preview,
                "metadata": {
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": df.columns.tolist(),
                    "sheet_names": pd.ExcelFile(file_path).sheet_names,
                },
            }
        except Exception as e:
            logger.error(f"XLSX parsing error: {str(e)}")
            raise

    async def _parse_ods(self, file_path: Path) -> Dict[str, Any]:
        """Parse ODS files."""
        try:
            df = pd.read_excel(file_path, engine="odf")
            preview = df.head().to_string()
            return {
                "content": df.to_dict(),
                "preview": preview,
                "metadata": {
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": df.columns.tolist(),
                },
            }
        except Exception as e:
            logger.error(f"ODS parsing error: {str(e)}")
            raise


class CodeParser(BaseParser):
    """Parser for code and configuration files (PY, JS, JSON, YAML, XML, TOML)."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        """Parse code/config file and return its contents and metadata.

        Args:
            file_path: Path to the code/config file

        Returns:
            Dict containing parsed content and metadata

        Raises:
            ValueError: If the file format is not supported
        """
        suffix = file_path.suffix.lower()
        try:
            if suffix == ".py":
                return await self._parse_python(file_path)
            elif suffix == ".js":
                return await self._parse_javascript(file_path)
            elif suffix == ".json":
                return await self._parse_json(file_path)
            elif suffix in [".yaml", ".yml"]:
                return await self._parse_yaml(file_path)
            elif suffix == ".xml":
                return await self._parse_xml(file_path)
            elif suffix == ".toml":
                return await self._parse_toml(file_path)
            else:
                raise ValueError(f"Unsupported code/config format: {suffix}")
        except Exception as e:
            logger.error(f"Error parsing code/config file {file_path}: {str(e)}")
            raise

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        """Generate a preview of the code/config content.

        Args:
            parsed_data: The parsed code/config data
            max_length: Maximum length of the preview

        Returns:
            A string preview of the content
        """
        preview = parsed_data.get("preview", "")
        if len(preview) > max_length:
            return preview[:max_length] + "..."
        return preview

    async def _parse_python(self, file_path: Path) -> Dict[str, Any]:
        """Parse Python files with AST analysis."""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                content = await f.read()

            tree = ast.parse(content)

            # Extract classes and functions
            classes = [
                node.name for node in ast.walk(tree) if isinstance(node, ast.ClassDef)
            ]
            functions = [
                node.name
                for node in ast.walk(tree)
                if isinstance(node, ast.FunctionDef)
            ]

            # Count imports
            imports = [
                node
                for node in ast.walk(tree)
                if isinstance(node, (ast.Import, ast.ImportFrom))
            ]

            return {
                "content": content,
                "preview": content[:1000] if len(content) > 1000 else content,
                "metadata": {
                    "classes": classes,
                    "functions": functions,
                    "import_count": len(imports),
                    "loc": len(content.splitlines()),
                },
            }
        except Exception as e:
            logger.error(f"Python parsing error: {str(e)}")
            raise

    async def _parse_javascript(self, file_path: Path) -> Dict[str, Any]:
        """Parse JavaScript files."""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                content = await f.read()

            return {
                "content": content,
                "preview": content[:1000] if len(content) > 1000 else content,
                "metadata": {"loc": len(content.splitlines())},
            }
        except Exception as e:
            logger.error(f"JavaScript parsing error: {str(e)}")
            raise

    async def _parse_json(self, file_path: Path) -> Dict[str, Any]:
        """Parse JSON files."""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                content = await f.read()
                data = json.loads(content)

            return {
                "content": data,
                "preview": content[:1000] if len(content) > 1000 else content,
                "metadata": {
                    "keys": list(data.keys()) if isinstance(data, dict) else None,
                    "size": len(data) if isinstance(data, (dict, list)) else None,
                },
            }
        except Exception as e:
            logger.error(f"JSON parsing error: {str(e)}")
            raise

    async def _parse_yaml(self, file_path: Path) -> Dict[str, Any]:
        """Parse YAML files."""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                content = await f.read()
                data = yaml.safe_load(content)

            return {
                "content": data,
                "preview": content[:1000] if len(content) > 1000 else content,
                "metadata": {
                    "keys": list(data.keys()) if isinstance(data, dict) else None,
                    "size": len(data) if isinstance(data, (dict, list)) else None,
                },
            }
        except Exception as e:
            logger.error(f"YAML parsing error: {str(e)}")
            raise

    async def _parse_xml(self, file_path: Path) -> Dict[str, Any]:
        """Parse XML files."""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()

            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                content = await f.read()

            return {
                "content": content,
                "preview": content[:1000] if len(content) > 1000 else content,
                "metadata": {
                    "root_tag": root.tag,
                    "children_count": len(root),
                    "attributes": dict(root.attrib),
                },
            }
        except Exception as e:
            logger.error(f"XML parsing error: {str(e)}")
            raise

    async def _parse_toml(self, file_path: Path) -> Dict[str, Any]:
        """Parse TOML files."""
        try:
            async with aiofiles.open(file_path, "rb") as f:
                content = await f.read()
                data = tomli.loads(content.decode("utf-8"))

            return {
                "content": data,
                "preview": str(data)[:1000],
                "metadata": {
                    "keys": list(data.keys()) if isinstance(data, dict) else None,
                    "size": len(data) if isinstance(data, (dict, list)) else None,
                },
            }
        except Exception as e:
            logger.error(f"TOML parsing error: {str(e)}")
            raise


class VectorParser(BaseParser):
    """Parser for CAD and vector files (DWG, DXF, SVG)."""

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        """Parse vector/CAD file and return its contents and metadata.

        Args:
            file_path: Path to the vector/CAD file

        Returns:
            Dict containing parsed content and metadata

        Raises:
            ValueError: If the file format is not supported
        """
        suffix = file_path.suffix.lower()
        try:
            if suffix in [".dwg", ".dxf"]:
                return await self._parse_cad(file_path)
            elif suffix == ".svg":
                return await self._parse_svg(file_path)
            else:
                raise ValueError(f"Unsupported vector format: {suffix}")
        except Exception as e:
            logger.error(f"Error parsing vector file {file_path}: {str(e)}")
            raise

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        """Generate a preview of the vector/CAD content.

        Args:
            parsed_data: The parsed vector/CAD data
            max_length: Maximum length of the preview

        Returns:
            A string preview of the content
        """
        preview = parsed_data.get("preview", "")
        if len(preview) > max_length:
            return preview[:max_length] + "..."
        return preview

    async def _parse_cad(self, file_path: Path) -> Dict[str, Any]:
        """Parse DWG/DXF files."""
        try:
            doc = ezdxf.readfile(str(file_path))
            modelspace = doc.modelspace()

            # Extract entities
            entities = {
                "lines": len(modelspace.query("LINE")),
                "circles": len(modelspace.query("CIRCLE")),
                "arcs": len(modelspace.query("ARC")),
                "polylines": len(modelspace.query("LWPOLYLINE")),
                "text": len(modelspace.query("TEXT")),
            }

            # Extract layers
            layers = [layer.dxf.name for layer in doc.layers]

            # Get document metadata
            metadata = {
                "filename": file_path.name,
                "created_by": doc.header["$TDCREATE"],
                "last_modified": doc.header["$TDUPDATE"],
                "drawing_units": doc.header["$INSUNITS"],
                "layers": layers,
                "entity_counts": entities,
            }

            return {
                "content": str(doc.entitydb),
                "preview": f"CAD drawing with {sum(entities.values())} entities across {len(layers)} layers",
                "metadata": metadata,
            }
        except Exception as e:
            logger.error(f"CAD parsing error: {str(e)}")
            raise

    async def _parse_svg(self, file_path: Path) -> Dict[str, Any]:
        """Parse SVG files."""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()

            width = root.get("width", "unknown")
            height = root.get("height", "unknown")
            viewBox = root.get("viewBox", "unknown")

            elements = {
                "path": len(root.findall(".//{*}path")),
                "rect": len(root.findall(".//{*}rect")),
                "circle": len(root.findall(".//{*}circle")),
                "text": len(root.findall(".//{*}text")),
                "group": len(root.findall(".//{*}g")),
            }

            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                content = await f.read()

            return {
                "content": content,
                "preview": f"SVG image ({width}x{height}) with {sum(elements.values())} elements",
                "metadata": {
                    "dimensions": {
                        "width": width,
                        "height": height,
                        "viewBox": viewBox,
                    },
                    "element_counts": elements,
                    "xmlns": root.get("xmlns", "unknown"),
                    "version": root.get("version", "unknown"),
                },
            }
        except Exception as e:
            logger.error(f"SVG parsing error: {str(e)}")
            raise


class ArchiveParser(BaseParser):
    """Parser for archive files (ZIP, TAR, RAR, 7Z).

    This parser extracts a list of contained files from the archive.
    """

    async def parse(self, file_path: Path) -> Dict[str, Any]:
        """Parse the archive file and return its contents and metadata.

        Args:
            file_path (Path): The path to the archive file.

        Returns:
            Dict[str, Any]: A dictionary containing the list of file names and count.

        Raises:
            Exception: If parsing fails or the archive format is unsupported.
        """
        try:
            suffix = file_path.suffix.lower()
            files_list = []
            if suffix == ".zip":
                with zipfile.ZipFile(file_path, "r") as archive:
                    files_list = archive.namelist()
            elif suffix in [".tar", ".tgz", ".gz", ".bz2"]:
                with tarfile.open(file_path, "r") as archive:
                    files_list = archive.getnames()
            elif suffix == ".rar":
                try:
                    with rarfile.RarFile(file_path, "r") as archive:
                        files_list = archive.namelist()
                except Exception as e:
                    if isinstance(e, rarfile.NotRarFile):
                        raise ValueError(f"Unsupported archive format: {suffix}") from e
                    else:
                        raise
            elif suffix == ".7z":
                with py7zr.SevenZipFile(file_path, "r") as archive:
                    files_list = archive.getnames()
            else:
                raise ValueError(f"Unsupported archive format: {suffix}")
            return {"files": files_list, "count": len(files_list)}
        except Exception as e:
            logger.error(f"Archive parsing failed for {file_path}: {e}", exc_info=True)
            raise

    def get_preview(self, parsed_data: Dict[str, Any], max_length: int = 1000) -> str:
        """Generate a preview of the archive contents.

        Args:
            parsed_data (Dict[str, Any]): The parsed archive data.
            max_length (int): Maximum length of the preview text.

        Returns:
            str: A string preview of the archive file names.
        """
        files = parsed_data.get("files", [])
        preview = "\n".join(files)
        if len(preview) > max_length:
            preview = preview[:max_length] + "..."
        return preview


def get_parser_for_mime_type(mime_type: str) -> Optional[BaseParser]:
    """
    Factory function to get the appropriate parser for a MIME type.

    Args:
        mime_type: The MIME type of the file

    Returns:
        An instance of the appropriate parser, or None if no parser is available
    """
    parsers = {
        "text/plain": TextParser(),
        "text/markdown": TextParser(),
        "text/html": TextParser(),
        "application/json": JsonParser(),
        "application/yaml": YamlParser(),
        "text/yaml": YamlParser(),
        "text/csv": CsvParser(),
        "application/xml": XmlParser(),
        "text/xml": XmlParser(),
        "application/pdf": DocumentParser(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DocumentParser(),
        "application/vnd.oasis.opendocument.text": DocumentParser(),
        "application/epub+zip": DocumentParser(),
        "application/zip": ArchiveParser(),
        "application/x-tar": ArchiveParser(),
        "application/x-rar-compressed": ArchiveParser(),
        "application/x-7z-compressed": ArchiveParser(),
        "application/gzip": ArchiveParser(),
        "application/x-bzip2": ArchiveParser(),
        "image/vnd.dxf": VectorParser(),
        "image/x-dwg": VectorParser(),
        "image/svg+xml": VectorParser(),
        "text/x-python": CodeParser(),
        "application/javascript": CodeParser(),
        "application/toml": CodeParser(),
    }

    if mime_type in parsers:
        return parsers[mime_type]

    main_type = mime_type.split("/")[0]
    if main_type == "text":
        return TextParser()

    return None
